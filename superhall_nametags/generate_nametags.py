import sys,math
from MCR_Dining.getnames import AttendeeScraper
from docx import Document




folder='/mnt/c/Users/Cole/Downloads'
folder='/home/colehunt/software/MCR-dining/data'
folder='/home/ach221/Downloads'
### Get the names from Upay and seating form responses to generate the Matrices required
event_booking_html = f"{folder}/Upay - Event Booking.html"
swaps_xls = f"{folder}/MTSuperhallSwaps2025-26.xlsx"
nametag_template=f"superhall_nametags.docx"
nametag_template=f"superhall_nametags_xmas.docx"
outname=f"nametags_filled"

### Get the names from Upay
guestlist=AttendeeScraper(event_booking_html,swaps_xls)
guestlist.load_Upay()
guestlist.load_Swaps()
everyone=guestlist.everyone
print(everyone)
### Clean up the numbering from the guestlist
everyone=[name.strip('(1)').strip('(2)').strip('(3)') for name in everyone]
ntot=len(everyone)

### Generate the nametags
doc = Document(nametag_template)
# get the number of placeholdes
table = doc.tables[0]
placeholder_count = 0
for row in table.rows:
    cell = row.cells[0]
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if "Guest Name" in run.text:
                placeholder_count += 1

                
n_outfiles=result = math.ceil(ntot / placeholder_count)
if n_outfiles == 1: outnames=[F'{outname}.docx']
else: outnames=[f'{outname}({i+1}).docx' for i in range(n_outfiles)]


ntot = len(everyone) # Number of names in the attendees list
name_counter=0 #counter for which person we using
for outname in outnames:
    # Insert names into the document
    doc = Document(nametag_template) #reload

    # Document is formatted as a table (see printdocstructure.py for more details)
    table=doc.tables[0] 
    for row_idx, row in enumerate(table.rows): # go through the cells in the docx
        if name_counter>=ntot: break #if we are at the end of the everyone list, then stop
        everyone[name_counter] = everyone[name_counter].strip() + "\n"
        cell = row.cells[0]
        for paragraph in cell.paragraphs:
            # Iterate through each run in the paragraph
            for run in paragraph.runs:
                # If "guest name" is found in the run text, replace it
                if "Guest Name" in run.text:  
                    if name_counter>=ntot:
                        # remove the  paragraph if there are no more names
                        p = paragraph._element
                        p.getparent().remove(p)
                        continue
                    # Replace 'guest name' with the current name in the attendees list
                    run.text = run.text.replace("Guest Name", everyone[name_counter])
                    name_counter+=1 #add one to the counter

    doc.save(outname)
    print(f"Saved to {outname}")    