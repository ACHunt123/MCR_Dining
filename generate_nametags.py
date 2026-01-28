import sys,math
from pathlib import Path
from MCR_Dining.getnames import AttendeeScraper
from docx import Document



folder='/home/colehunt/software/MCR-dining/data'
folder='/home/ach221/Desktop'
### Input files
event_booking_html = f"{folder}/Upay - Event Booking.html"
swaps_xls = f"{folder}/MTSuperhallSwaps2025-26.xlsx"

### Nametag Template
nametag_template=["superhall_nametags.docx","superhall_nametags_xmas.docx"][0]
outname=f"nametags_filled"

### Get the names from Upay
guestlist=AttendeeScraper()
guestlist.load_Upay(event_booking_html)
# guestlist.load_Swaps(swaps_xls)

### Make sure no duplications and the correct number of people
ntot=len(guestlist.everyone)
duplicates=len(guestlist.everyone)==len(list(set(guestlist.everyone)))
print(f"there are {['','no'][duplicates]} duplicates in the list of names")
print(f'there are {ntot} people total (including dup.s if present)')

### Clean up the numbering from the guestlist
everyone=[name.strip('(1)').strip('(2)').strip('(3)') for name in guestlist.everyone]


### Generate the nametags
loc=Path(__file__).parent # Get the path to the current script
nametag_template_loc=loc/"superhall_nametags"/nametag_template
doc = Document(nametag_template_loc)
# get the number of placeholdes
table = doc.tables[0]
placeholder_count = 0
for row in table.rows:
    cell = row.cells[0]
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if "Guest Name" in run.text:
                placeholder_count += 1

                
n_outfiles=result = math.ceil(ntot / placeholder_count)
if n_outfiles == 1: outnames=[F'{outname}.docx']
else: outnames=[f'{outname}({i+1}).docx' for i in range(n_outfiles)]


name_counter=0 #counter for which person we using
for outname in outnames:
    # Insert names into the document
    doc = Document(nametag_template_loc) #reload

    # Document is formatted as a table (see printdocstructure.py for more details)
    table=doc.tables[0] 
    for row_idx, row in enumerate(table.rows): # go through the cells in the docx
        if name_counter>=ntot: break #if we are at the end of the everyone list, then stop
        everyone[name_counter] = everyone[name_counter].strip() + "\n"
        cell = row.cells[0]
        for paragraph in cell.paragraphs:
            # Iterate through each run in the paragraph
            for run in paragraph.runs:
                # If "guest name" is found in the run text, replace it
                if "Guest Name" in run.text:  
                    if name_counter>=ntot:
                        # remove the  paragraph if there are no more names
                        p = paragraph._element
                        p.getparent().remove(p)
                        continue
                    # Replace 'guest name' with the current name in the attendees list
                    run.text = run.text.replace("Guest Name", everyone[name_counter])
                    name_counter+=1 #add one to the counter

    doc.save(outname)
    print(f"Saved to {outname}")    